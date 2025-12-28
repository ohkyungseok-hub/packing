import io
import json
import re
from datetime import date
from typing import Dict, List, Optional, Tuple

import streamlit as st
from docx import Document


def norm(s: str) -> str:
    # 필요하면 하이픈/공백 제거 규칙을 통일할 수 있어요.
    return str(s).strip()


def extract_ship_id(doc: Document) -> Optional[str]:
    """
    문서 전체에서 '출고주문번호:' 다음 값을 찾는다.
    예: '출고주문번호: OUT-0001'
    """
    pattern = re.compile(r"출고주문번호[:\s]*([A-Za-z0-9\-]+)")
    for p in doc.paragraphs:
        m = pattern.search(p.text)
        if m:
            return norm(m.group(1))
    return None


def parse_items_from_tables(doc: Document) -> Tuple[List[Dict], List[str]]:
    """
    표 헤더에서 '주문상품', '주문수량', '상품 바코드' 컬럼을 찾아
    각 행을 items로 변환한다.
    반환:
      items: [{barcode, name, qty}, ...]
      warnings: 경고 메시지 리스트
    """
    items: List[Dict] = []
    warnings: List[str] = []

    wanted = ["주문상품", "주문수량", "상품 바코드"]

    found_any_table = False
    for table in doc.tables:
        if not table.rows:
            continue

        headers = [norm(c.text) for c in table.rows[0].cells]
        if not all(w in headers for w in wanted):
            continue

        found_any_table = True
        idx_name = headers.index("주문상품")
        idx_qty = headers.index("주문수량")
        idx_barcode = headers.index("상품 바코드")

        for r_i, row in enumerate(table.rows[1:], start=2):  # 1-based row number 느낌
            cells = [norm(c.text) for c in row.cells]
            if len(cells) <= max(idx_name, idx_qty, idx_barcode):
                continue

            name = norm(cells[idx_name])
            barcode = norm(cells[idx_barcode])

            # qty 파싱 (워드 표에 이상한 공백/줄바꿈이 있을 수 있어서 숫자만 추출)
            qty_raw = norm(cells[idx_qty])
            m = re.search(r"\d+", qty_raw)
            qty = int(m.group(0)) if m else 0

            # 합계 행/빈 행 방지
            if not name and not barcode:
                continue
            if "합계" in name:
                continue

            if qty <= 0:
                warnings.append(f"표 {r_i}행: 수량 파싱 실패(값='{qty_raw}') → 건너뜀")
                continue

            if not barcode:
                warnings.append(f"표 {r_i}행: 상품 바코드가 비어있음(상품='{name}')")
                # 바코드 없는 건은 검증 불가라서 기본은 제외
                continue

            items.append({
                "barcode": barcode,
                "name": name if name else "(상품명 없음)",
                "qty": qty
            })

    if not found_any_table:
        warnings.append("필요한 헤더(주문상품/주문수량/상품 바코드)가 있는 표를 찾지 못했습니다.")

    return items, warnings


def build_orders_json(doc: Document) -> Tuple[Dict, List[str]]:
    warnings: List[str] = []

    ship_id = extract_ship_id(doc)
    if not ship_id:
        warnings.append("문서에서 '출고주문번호'를 찾지 못했습니다. (예: '출고주문번호: OUT-0001')")
        ship_id = "UNKNOWN"

    items, w2 = parse_items_from_tables(doc)
    warnings.extend(w2)

    orders_file = {
        "date": str(date.today()),
        "orders": [
            {
                "orderId": ship_id,
                "items": items
            }
        ]
    }

    if not items:
        warnings.append("추출된 상품(item)이 0건입니다. 워드 표 구조/헤더명을 확인하세요.")

    return orders_file, warnings


st.set_page_config(page_title="DOCX → orders.json 변환", layout="centered")
st.title("워드(.docx) → orders.json 변환기 (오프라인 피킹앱용)")

st.write(
    "1) 워드(.docx) 파일을 업로드하면\n"
    "2) 출고번호(orderId) + 상품바코드/수량을 추출해서\n"
    "3) `orders.json` 파일로 만들어 다운로드할 수 있습니다."
)

uploaded = st.file_uploader("워드(.docx) 업로드", type=["docx"])

if uploaded is not None:
    try:
        data = uploaded.read()
        doc = Document(io.BytesIO(data))

        orders_file, warnings = build_orders_json(doc)

        ship_id = orders_file["orders"][0]["orderId"]
        items = orders_file["orders"][0]["items"]

        st.success("변환 성공!")
        st.subheader("추출 결과 요약")
        st.write(f"- 출고번호(orderId): **{ship_id}**")
        st.write(f"- 상품(item) 건수: **{len(items)}**")

        if warnings:
            st.warning("경고/확인 필요")
            for w in warnings:
                st.write(f"- {w}")

        if items:
            st.subheader("상품 리스트(미리보기)")
            st.dataframe(items, use_container_width=True)

        st.subheader("orders.json 미리보기")
        st.code(json.dumps(orders_file, ensure_ascii=False, indent=2), language="json")

        # 다운로드 버튼
        out_bytes = json.dumps(orders_file, ensure_ascii=False, indent=2).encode("utf-8")
        st.download_button(
            label="orders.json 다운로드",
            data=out_bytes,
            file_name="orders.json",
            mime="application/json"
        )

        st.info("다운로드한 orders.json을 피킹앱에서 'orders.json 업로드' 버튼으로 넣으면 됩니다.")

    except Exception as e:
        st.error("변환 중 오류가 발생했습니다.")
        st.code(str(e))
        st.write("✅ 확인사항:")
        st.write("- 업로드한 파일이 진짜 `.docx`인지 (워드에서 '다른 이름으로 저장'으로 docx로 저장)")
        st.write("- 문서에 '출고주문번호:' 문구가 있는지")
        st.write("- 표 헤더에 '주문상품', '주문수량', '상품 바코드'가 정확히 있는지")
else:
    st.caption("워드 파일을 업로드하면 여기서 변환이 시작됩니다.")
